"""
OPTION B: Complete Fix for JBHI_FINAL_SUBMISSION_CORRECTED.docx
Fixes:
1. Wilson CI: [95.3%, 99.1%] → [95.2%, 99.0%] (Abstract + Results IV.A)
2. Add "400 windows combined" mention in Results summary
3. Update Conclusion with specific numbers (97.8% VitalDB, 25.4% CHARIS)
4. Note: CHARIS section already exists as Section IV.F - no need to add IV.E
"""

from docx import Document
import copy

input_path = '/home/ubuntu/upload/JBHI_FINAL_SUBMISSION_CORRECTED.docx'
output_path = '/home/ubuntu/upload/JBHI_FINAL_SUBMISSION_FINAL.docx'

doc = Document(input_path)

fixes_applied = []

for i, para in enumerate(doc.paragraphs):
    
    # FIX 1: Wilson CI in Abstract and Results
    # [95.3%, 99.1%] → [95.2%, 99.0%]
    if '95.3%' in para.text:
        for run in para.runs:
            if '95.3%' in run.text:
                run.text = run.text.replace('95.3%', '95.2%')
                fixes_applied.append(f"FIX 1a: Changed Wilson CI lower '95.3%' → '95.2%' (para {i})")
    
    if '99.1%' in para.text:
        for run in para.runs:
            if '99.1%' in run.text:
                run.text = run.text.replace('99.1%', '99.0%')
                fixes_applied.append(f"FIX 1b: Changed Wilson CI upper '99.1%' → '99.0%' (para {i})")
    
    # FIX 2: Update Conclusion with specific numbers
    # Find the sentence about "severe signal-level violations are pervasive"
    if 'severe signal-level violations are pervasive' in para.text.lower() or \
       'severe signal-\nlevel violations are pervasive' in para.text.lower():
        for run in para.runs:
            if 'severe signal-level violations are pervasive' in run.text.lower() or \
               'severe signal-' in run.text.lower():
                old_text = run.text
                # Don't modify here, we'll handle conclusion differently
                pass
    
    # FIX 3: In Conclusion - add specific numbers after "conservative integrity criteria"
    # Look for the conclusion text about "pervasive under conservative integrity criteria"
    if 'pervasive under conservative integrity criteria' in para.text:
        for run in para.runs:
            if 'conservative integrity criteria' in run.text:
                run.text = run.text.replace(
                    'conservative integrity criteria',
                    'conservative integrity criteria—with 97.8% of VitalDB windows and 25.4% of CHARIS windows failing at least one criterion (297 of 400 total windows across both databases)'
                )
                fixes_applied.append(f"FIX 3: Added specific numbers to Conclusion (para {i})")
                break
    
    # FIX 4: In the Summary of Key Findings (Section IV.H) - add combined total
    if 'The principal findings of this study can be summarized' in para.text:
        # Mark this paragraph for reference
        fixes_applied.append(f"FOUND: Summary of Key Findings at para {i}")

# Also check for Wilson CI in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if '95.3%' in para.text:
                    for run in para.runs:
                        if '95.3%' in run.text:
                            run.text = run.text.replace('95.3%', '95.2%')
                            fixes_applied.append("FIX 1c: Changed Wilson CI '95.3%' → '95.2%' in Table")
                if '99.1%' in para.text:
                    for run in para.runs:
                        if '99.1%' in run.text:
                            run.text = run.text.replace('99.1%', '99.0%')
                            fixes_applied.append("FIX 1d: Changed Wilson CI '99.1%' → '99.0%' in Table")

# Now find the Summary section and add combined total
for i, para in enumerate(doc.paragraphs):
    # Add combined total after the CHARIS findings in Section IV.H
    if 'signal integrity violations generalize beyond a single dataset' in para.text and \
       '25.4%' in para.text:
        for run in para.runs:
            if 'clinical context' in run.text or 'environment' in run.text:
                if 'Across both databases' not in para.text:
                    run.text = run.text.rstrip()
                    if run.text.endswith('.'):
                        run.text = run.text + ' Across both databases, 297 of 400 analyzed windows (74.2%) failed at least one integrity criterion, underscoring the widespread nature of ABP signal quality issues in open clinical repositories.'
                    fixes_applied.append(f"FIX 4: Added combined 400 windows total to Summary (para {i})")
                break

doc.save(output_path)

print("="*70)
print("  OPTION B FIXES APPLIED")
print("="*70)
for fix in fixes_applied:
    print(f"  {fix}")
print(f"\nTotal fixes: {len(fixes_applied)}")
print(f"Saved to: {output_path}")
print("="*70)
